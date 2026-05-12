from __future__ import annotations

import logging
import os
from io import BytesIO

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """文档加载器，支持PDF、Word和TXT格式"""

    @staticmethod
    def load_pdf(file_path: str) -> str:
        """加载PDF文档"""
        try:
            reader = PdfReader(file_path)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            return text
        except Exception as e:
            logger.error(f'Error loading PDF file {file_path}: {str(e)}')
            raise

    @staticmethod
    def load_pdf_from_bytes(file_bytes: bytes) -> str:
        """从字节流加载PDF文档"""
        try:
            reader = PdfReader(BytesIO(file_bytes))
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            return text
        except Exception as e:
            logger.error(f'Error loading PDF from bytes: {str(e)}')
            raise

    @staticmethod
    def load_word(file_path: str) -> str:
        """加载Word文档"""
        try:
            doc = Document(file_path)
            text = ''
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            return text
        except Exception as e:
            logger.error(f'Error loading Word file {file_path}: {str(e)}')
            raise

    @staticmethod
    def load_word_from_bytes(file_bytes: bytes) -> str:
        """从字节流加载Word文档"""
        try:
            doc = Document(BytesIO(file_bytes))
            text = ''
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            return text
        except Exception as e:
            logger.error(f'Error loading Word from bytes: {str(e)}')
            raise

    @staticmethod
    def load_txt(file_path: str) -> str:
        """加载TXT文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f'Error loading TXT file {file_path}: {str(e)}')
            raise

    @staticmethod
    def load_txt_from_bytes(file_bytes: bytes) -> str:
        """从字节流加载TXT文档"""
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError as e:
            logger.error(f'Error decoding TXT from bytes: {str(e)}')
            raise ValueError(f'File encoding must be UTF-8: {str(e)}')

    @staticmethod
    def load_document(file_path: str) -> str:
        """根据文件扩展名加载相应类型的文档"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f'File not found: {file_path}')

        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == '.pdf':
            return DocumentLoader.load_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentLoader.load_word(file_path)
        elif ext == '.txt':
            return DocumentLoader.load_txt(file_path)
        else:
            raise ValueError(f'Unsupported file format: {ext}')

    @staticmethod
    def load_document_from_bytes(file_bytes: bytes, filename: str) -> str:
        """根据文件名从字节流加载文档

        Args:
            file_bytes: 文件字节流
            filename: 文件名（用于判断扩展名）

        Returns:
            文档文本内容
        """
        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        if ext == '.pdf':
            return DocumentLoader.load_pdf_from_bytes(file_bytes)
        elif ext in ['.docx', '.doc']:
            return DocumentLoader.load_word_from_bytes(file_bytes)
        elif ext in ['.txt', '.md', '.json', '.yaml', '.yml']:
            return DocumentLoader.load_txt_from_bytes(file_bytes)
        else:
            raise ValueError(
                f'Unsupported file format: {ext}. '
                'Supported: .pdf, .docx, .doc, .txt, .md, .json, .yaml, .yml'
            )